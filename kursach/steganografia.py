# Стеганография в DOCX документах
from tkinter import *
from tkinter import filedialog, messagebox, scrolledtext, ttk
from docx import Document
from docx.shared import Pt
import math
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg


class SteganographyApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Стеганография в DOCX документах ")
        self.root.geometry("900x850")
        self.current_document = None
        self.current_text = ""
        self.current_file_path = ""
        self.hidden_message_positions = []
        self.base_font_size = 11
        self.num_fonts = 32
        self.full_font_mapping = {
            'А': ('Cambria', 13),
            'Б': ('Helvetica', 11),
            'В': ('Verdana', 11),
            'Г': ('Tahoma', 11),
            'Д': ('Trebuchet MS', 11),
            'Е': ('Calibri', 12),
            'Ё': ('Candara', 12),
            'Ж': ('Segoe UI', 11),
            'З': ('Roboto', 11),
            'И': ('Open Sans', 11),
            'Й': ('Lato', 11),
            'К': ('Montserrat', 11),
            'Л': ('Constantia', 12),
            'М': ('Ubuntu', 11),
            'Н': ('Noto Sans', 11),
            'О': ('PT Sans', 11),
            'П': ('Fira Sans', 11),
            'Р': ('Droid Sans', 12),
            'С': ('Nunito', 12),
            'Т': ('Charter', 12),
            'У': ('Quicksand', 12),
            'Ф': ('Inter', 12),
            'Х': ('Manrope', 12),
            'Ц': ('Mulish', 12),
            'Ч': ('Outfit', 12),
            'Ш': ('Baskerville', 12),
            'Щ': ('Rubik', 12),
            'Ъ': ('Golos Text', 12),
            'Ы': ('Literata', 12),
            'Ь': ('IBM Plex Sans', 12),
            'Э': ('Franklin Gothic', 12),
            'Ю': ('Lucida Sans', 12),
            'Я': ('Commissioner', 12)
        }
        self.all_fonts = []
        seen_fonts = set()
        for letter, (font, size) in self.full_font_mapping.items():
            if font not in seen_fonts:
                seen_fonts.add(font)
                self.all_fonts.append((font, size))
        self.all_fonts.sort(key=lambda x: x[0])
        self.current_fonts = self.all_fonts[:self.num_fonts]
        self.font_index_mapping = {i: font for i, font in enumerate(self.current_fonts)}
        self.reverse_font_mapping = {}
        for idx, (font, size) in enumerate(self.current_fonts):
            key = f"{font}_{int(size)}"
            self.reverse_font_mapping[key] = idx
        self.encoding_step = 1
        self.message_entry = None
        self.create_widgets()
        if self.message_entry:
            self.message_entry.bind('<KeyRelease>', self.on_text_change)
        self.update_capacity_display()

    def text_to_bits(self, text):
        if not text:
            return ""
        byte_data = text.encode('utf-8')
        bits = ''.join(format(byte, '08b') for byte in byte_data)
        return bits

    def bits_to_text(self, bits):
        if not bits or len(bits) % 8 != 0:
            return ""
        byte_data = bytes(int(bits[i:i + 8], 2) for i in range(0, len(bits), 8))
        try:
            text = byte_data.decode('utf-8')
            return text
        except UnicodeDecodeError:
            return ""

    def bits_to_chunks(self, bits, k):
        chunks = []
        padding_bits = 0
        for i in range(0, len(bits), k):
            chunk = bits[i:i + k]
            if len(chunk) < k:
                padding_bits = k - len(chunk)
                chunk = chunk + '0' * padding_bits
            chunks.append(chunk)
        return chunks, padding_bits

    def chunk_to_index(self, chunk):
        if not chunk:
            return 0
        return int(chunk, 2)

    def index_to_bits(self, index, k):
        return format(index, '0{}b'.format(k))

    def update_font_subset(self):
        self.current_fonts = self.all_fonts[:self.num_fonts]
        self.font_index_mapping = {i: font for i, font in enumerate(self.current_fonts)}
        self.reverse_font_mapping = {}
        for idx, (font, size) in enumerate(self.current_fonts):
            key = f"{font}_{int(size)}"
            self.reverse_font_mapping[key] = idx

    def get_capacity(self):
        if self.num_fonts == 0:
            return 0
        return math.log2(self.num_fonts)

    def get_k_bits(self):
        return int(self.get_capacity())

    def get_max_capacity(self, text_length=None):
        if text_length is None:
            if not self.current_text:
                return 0, 0
            text_length = self.count_russian_letters(self.current_text)
        available_positions = self.get_available_positions_count(text_length)
        k = self.get_k_bits()
        max_capacity_bits = available_positions * k
        if k > 0:
            max_capacity_chars = (max_capacity_bits - 16) // 8 if max_capacity_bits > 16 else 0
        else:
            max_capacity_chars = 0
        return max_capacity_chars, max_capacity_bits

    def update_capacity_display(self):
        capacity = self.get_capacity()
        k = self.get_k_bits()
        if hasattr(self, 'capacity_label'):
            self.capacity_label.config(
                text=f"Ёмкость канала: {capacity:.2f} бит/символ (N={self.num_fonts}, k={k} бит)")
        if self.current_document:
            self.update_statistics()

    def on_font_count_change(self, event=None):
        selected = self.font_count_combo.get()
        self.num_fonts = int(selected.split()[0])
        self.update_font_subset()
        self.update_capacity_display()
        self.update_capacity_indicator()
        messagebox.showinfo("Информация",
                            f"Теперь используется {self.num_fonts} шрифтов.\n"
                            f"Ёмкость канала: {self.get_capacity():.2f} бит/символ (k={self.get_k_bits()} бит)\n\n"
                            f"Сообщения кодируются с помощью {self.num_fonts} различных\n"
                            f"шрифтов, что позволяет передавать до {self.get_k_bits()} бит на символ.")

    def create_widgets(self):
        top_frame = Frame(self.root)
        top_frame.pack(pady=5)
        self.open_button = Button(top_frame, text="Открыть DOCX",
                                  command=self.open_document,
                                  bg="#4CAF50", fg="white",
                                  padx=15, pady=5)
        self.open_button.pack(side=LEFT, padx=5)
        self.save_button = Button(top_frame, text="Сохранить DOCX",
                                  command=self.save_document,
                                  bg="#2196F3", fg="white",
                                  padx=15, pady=5)
        self.save_button.pack(side=LEFT, padx=5)
        self.load_cipher_button = Button(top_frame, text="Загрузить шифр-текст из DOCX",
                                         command=self.load_cipher_text,
                                         bg="#FF5722", fg="white",
                                         padx=15, pady=5)
        self.load_cipher_button.pack(side=LEFT, padx=5)
        fonts_frame = LabelFrame(self.root, text="Выбор ёмкости канала", padx=5, pady=5)
        fonts_frame.pack(fill=X, padx=10, pady=5)
        font_count_frame = Frame(fonts_frame)
        font_count_frame.pack(fill=X, pady=2)
        Label(font_count_frame, text="Количество используемых шрифтов (N):",
              font=("Arial", 10, "bold")).pack(side=LEFT, padx=5)
        self.font_count_var = StringVar(value="32 шрифта")
        self.font_count_combo = ttk.Combobox(font_count_frame, textvariable=self.font_count_var,
                                             values=["2 шрифта", "4 шрифта", "8 шрифтов",
                                                     "16 шрифтов", "32 шрифта"],
                                             state="readonly", width=15)
        self.font_count_combo.pack(side=LEFT, padx=5)
        self.font_count_combo.bind("<<ComboboxSelected>>", self.on_font_count_change)
        self.capacity_label = Label(fonts_frame, text="Ёмкость канала: 0.00 бит/символ",
                                    font=("Arial", 10, "bold"), fg="blue")
        self.capacity_label.pack(anchor=W, pady=2)
        self.plot_button = Button(fonts_frame, text="Построить график зависимости C = log2(N)",
                                  command=self.plot_capacity_graph,
                                  bg="#9C27B0", fg="white",
                                  padx=20, pady=3)
        self.plot_button.pack(anchor=W, pady=2)
        stats_frame = LabelFrame(self.root, text="Статистика документа", padx=5, pady=5)
        stats_frame.pack(fill=X, padx=10, pady=5)
        self.stats_letters = Label(stats_frame, text="Букв: 0", font=("Arial", 9, "bold"))
        self.stats_letters.pack(side=LEFT, padx=10)
        self.stats_words = Label(stats_frame, text="Слов: 0", font=("Arial", 9, "bold"))
        self.stats_words.pack(side=LEFT, padx=10)
        self.stats_capacity = Label(stats_frame, text="Вместимость: 0 символов (0 бит)",
                                    font=("Arial", 9, "bold"), fg="green")
        self.stats_capacity.pack(side=LEFT, padx=10)
        text_frame = LabelFrame(self.root, text="Текст документа", padx=5, pady=5)
        text_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        self.text_display = scrolledtext.ScrolledText(text_frame, wrap=WORD,
                                                      width=80, height=8,
                                                      font=("Arial", 11))
        self.text_display.pack(fill=BOTH, expand=True)
        message_frame = LabelFrame(self.root, text="Скрытое сообщение", padx=5, pady=5)
        message_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        self.message_entry = scrolledtext.ScrolledText(message_frame, wrap=WORD,
                                                       width=80, height=3,
                                                       font=("Arial", 11))
        self.message_entry.pack(fill=BOTH, expand=True)
        self.capacity_indicator = Label(message_frame, text="Текущее сообщение: 0/0 символов (0%)",
                                        font=("Arial", 8), fg="blue")
        self.capacity_indicator.pack(anchor=W, pady=2)
        settings_frame = LabelFrame(self.root, text="Настройки кодирования", padx=5, pady=5)
        settings_frame.pack(fill=X, padx=10, pady=5)
        step_frame = Frame(settings_frame)
        step_frame.pack(fill=X, pady=2)
        Label(step_frame, text="Шаг кодирования (1-10):").pack(side=LEFT, padx=5)
        self.step_var = IntVar(value=1)
        step_spinbox = Spinbox(step_frame, from_=1, to=10, textvariable=self.step_var,
                               width=5, command=self.on_step_change)
        step_spinbox.pack(side=LEFT, padx=5)
        self.positions_label = Label(settings_frame, text="", font=("Arial", 8), fg="purple")
        self.positions_label.pack(anchor=W, pady=2)
        action_frame = Frame(self.root)
        action_frame.pack(pady=5)
        self.encrypt_button = Button(action_frame, text="Зашифровать",
                                     command=self.encrypt_message,
                                     bg="#FF9800", fg="white",
                                     padx=20, pady=5)
        self.encrypt_button.pack(side=LEFT, padx=5)
        self.decrypt_button = Button(action_frame, text="Расшифровать",
                                     command=self.decrypt_message,
                                     bg="#9C27B0", fg="white",
                                     padx=20, pady=5)
        self.decrypt_button.pack(side=LEFT, padx=5)
        self.clear_button = Button(action_frame, text="Очистить от шифра",
                                   command=self.clear_encryption,
                                   bg="#607D8B", fg="white",
                                   padx=20, pady=5)
        self.clear_button.pack(side=LEFT, padx=5)
        self.show_map_button = Button(action_frame, text="Показать карту скрытия",
                                      command=self.show_encryption_map,
                                      bg="#795548", fg="white",
                                      padx=20, pady=5)
        self.show_map_button.pack(side=LEFT, padx=5)
        result_frame = LabelFrame(self.root, text="Результат расшифровки", padx=5, pady=5)
        result_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        self.result_display = scrolledtext.ScrolledText(result_frame, wrap=WORD,
                                                        width=80, height=3,
                                                        font=("Arial", 11))
        self.result_display.pack(fill=BOTH, expand=True)
        self.save_result_button = Button(self.root, text="Сохранить расшифровку",
                                         command=self.save_decryption_result,
                                         bg="#f44336", fg="white",
                                         padx=20, pady=5)
        self.save_result_button.pack(pady=5)

    def plot_capacity_graph(self):
        graph_window = Toplevel(self.root)
        graph_window.title("График зависимости ёмкости канала")
        graph_window.geometry("800x600")
        N_values = [2, 4, 8, 16, 32]
        C_values = [math.log2(n) for n in N_values]
        fig, ax = plt.subplots(figsize=(8, 6), dpi=100)
        ax.plot(N_values, C_values, 'b-o', linewidth=2, markersize=8, label='C = log₂(N)')
        for i, (n, c) in enumerate(zip(N_values, C_values)):
            ax.annotate(f'({n}, {c:.2f})',
                        xy=(n, c),
                        xytext=(5, 5),
                        textcoords='offset points',
                        fontsize=10,
                        bbox=dict(boxstyle='round,pad=0.3', facecolor='yellow', alpha=0.7))
        ax.set_xlabel('Количество шрифтов (N)', fontsize=12, fontweight='bold')
        ax.set_ylabel('Ёмкость канала (бит/символ)', fontsize=12, fontweight='bold')
        ax.set_title('Зависимость ёмкости канала от количества используемых шрифтов\nC = log₂(N)',
                     fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.set_xlim(0, 35)
        ax.set_ylim(0, 5.5)
        ax.grid(True, alpha=0.3)
        ax.legend(loc='lower right', fontsize=10)
        ax.text(25, 4.5,
                f'C = log₂(N)\n\nПример:\nN=2 → C=1.00 бит/символ\nN=4 → C=2.00 бит/символ\nN=8 → C=3.00 бит/символ\nN=16 → C=4.00 бит/символ\nN=32 → C=5.00 бит/символ',
                bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.8),
                fontsize=9, verticalalignment='top')
        canvas = FigureCanvasTkAgg(fig, master=graph_window)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=BOTH, expand=True, padx=10, pady=10)
        Button(graph_window, text="Закрыть",
               command=graph_window.destroy,
               bg="#607D8B", fg="white",
               padx=20, pady=5).pack(pady=10)
        messagebox.showinfo("График построен",
                            f"График зависимости ёмкости канала от количества шрифтов построен.\n\n"
                            f"Текущее значение при N={self.num_fonts}: {self.get_capacity():.2f} бит/символ")

    def on_step_change(self):
        self.encoding_step = self.step_var.get()
        self.update_positions_info()
        self.update_capacity_indicator()

    def update_positions_info(self):
        if not self.current_document:
            return
        letters_count = self.count_russian_letters(self.current_text)
        step = self.encoding_step
        available = (letters_count + step - 1) // step
        info = f"Шаг {step}: будет использована каждая {step}-я буква (доступно: {available} позиций)"
        self.positions_label.config(text=info)

    def count_russian_letters(self, text):
        count = 0
        for char in text:
            if 'А' <= char.upper() <= 'Я' or char.upper() == 'Ё':
                count += 1
        return count

    def count_words(self, text):
        words = text.split()
        return len(words)

    def update_statistics(self):
        if not self.current_document:
            return
        text = self.text_display.get(1.0, END).strip()
        letters_count = self.count_russian_letters(text)
        words_count = self.count_words(text)
        max_chars, max_bits = self.get_max_capacity(letters_count)
        capacity_text = f"Вместимость: {max_chars} символов ({max_bits:.0f} бит)"
        if letters_count > 0:
            efficiency = (max_bits / letters_count) if letters_count > 0 else 0
            capacity_text += f" (эффективность: {efficiency:.2f} бит/букву)"
        self.stats_letters.config(text=f"Букв: {letters_count}")
        self.stats_words.config(text=f"Слов: {words_count}")
        self.stats_capacity.config(text=capacity_text)
        self.update_positions_info()
        self.update_capacity_indicator()

    def update_capacity_indicator(self):
        if not self.current_document:
            self.capacity_indicator.config(text="Текущее сообщение: 0/0 символов (0%)")
            return
        letters_count = self.count_russian_letters(self.current_text)
        if self.message_entry:
            message = self.message_entry.get(1.0, END).strip()
        else:
            message = ""
        message_length = len(message)
        max_chars, max_bits = self.get_max_capacity(letters_count)
        if message_length > max_chars:
            percent = 100
            color = "red"
            status = f"ПРЕВЫШЕНИЕ! {message_length} > {max_chars} символов"
        else:
            percent = (message_length / max_chars * 100) if max_chars > 0 else 0
            if percent < 50:
                color = "green"
            elif percent < 80:
                color = "orange"
            else:
                color = "red"
            status = f"{message_length}/{max_chars} символов ({percent:.1f}%)"
        self.capacity_indicator.config(
            text=f"Текущее сообщение: {status} | Ёмкость: {self.get_capacity():.2f} бит/символ",
            fg=color
        )

    def get_available_positions_count(self, total_letters):
        step = self.encoding_step
        return (total_letters + step - 1) // step

    def get_encoding_positions(self, total_letters, needed_positions):
        step = self.encoding_step
        positions = [i for i in range(0, total_letters, step)][:needed_positions]
        return positions

    def load_cipher_text(self):
        try:
            file_path = filedialog.askopenfilename(
                title="Выберите DOCX файл с текстом для шифрования",
                filetypes=[("DOCX файлы", "*.docx"), ("Все файлы", "*.*")]
            )
            if not file_path:
                return
            doc = Document(file_path)
            cipher_text = ""
            for paragraph in doc.paragraphs:
                cipher_text += paragraph.text + "\n"
            if self.message_entry:
                self.message_entry.delete(1.0, END)
                self.message_entry.insert(1.0, cipher_text.strip())
                self.update_capacity_indicator()
            messagebox.showinfo("Успех", "Текст успешно загружен из DOCX!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить файл: {str(e)}")

    def check_for_hidden_message(self):
        if not self.current_document:
            return False
        for paragraph in self.current_document.paragraphs:
            for run in paragraph.runs:
                font_name = run.font.name
                font_size = run.font.size.pt if run.font.size else None
                if font_name and font_size:
                    key = f"{font_name}_{int(font_size)}"
                    if key in self.reverse_font_mapping:
                        return True
        return False

    def open_document(self):
        try:
            file_path = filedialog.askopenfilename(
                title="Выберите DOCX файл",
                filetypes=[("DOCX файлы", "*.docx"), ("Все файлы", "*.*")]
            )
            if not file_path:
                return
            self.current_document = Document(file_path)
            self.current_file_path = file_path
            self.current_text = ""
            for paragraph in self.current_document.paragraphs:
                self.current_text += paragraph.text + "\n"
            self.text_display.delete(1.0, END)
            self.text_display.insert(1.0, self.current_text)
            self.update_statistics()
            if self.check_for_hidden_message():
                answer = messagebox.askyesno("Внимание",
                                             "Документ содержит скрытое сообщение!\n"
                                             "Хотите его расшифровать?")
                if answer:
                    self.decrypt_message()
            messagebox.showinfo("Успех", "Документ успешно загружен!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть файл: {str(e)}")

    def save_document(self):
        if not self.current_document:
            messagebox.showwarning("Предупреждение", "Сначала откройте документ!")
            return
        try:
            file_path = filedialog.asksaveasfilename(
                title="Сохранить документ",
                defaultextension=".docx",
                filetypes=[("DOCX файлы", "*.docx")]
            )
            if not file_path:
                return
            self.current_document.save(file_path)
            messagebox.showinfo("Успех", "Документ успешно сохранен!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")

    def encrypt_message(self):
        if not self.current_document:
            messagebox.showwarning("Предупреждение", "Сначала откройте документ!")
            return
        if self.check_for_hidden_message():
            answer = messagebox.askyesno("Предупреждение",
                                         "Документ уже содержит скрытое сообщение!\n"
                                         "Хотите перезаписать его новым?")
            if not answer:
                return
        if not self.message_entry:
            messagebox.showwarning("Предупреждение", "Поле для сообщения не найдено!")
            return
        hidden_message = self.message_entry.get(1.0, END).strip()
        if not hidden_message:
            messagebox.showwarning("Предупреждение", "Введите скрытое сообщение!")
            return
        doc_text = self.text_display.get(1.0, END).strip()
        letter_count = self.count_russian_letters(doc_text)
        max_chars, max_bits = self.get_max_capacity(letter_count)
        if len(hidden_message) > max_chars:
            messagebox.showerror("Ошибка",
                                 f"Скрытое сообщение слишком длинное!\n"
                                 f"Длина сообщения: {len(hidden_message)} символов\n"
                                 f"Максимальная вместимость: {max_chars} символов\n"
                                 f"Ёмкость канала: {self.get_capacity():.2f} бит/символ\n"
                                 f"k = {self.get_k_bits()} бит на символ")
            return
        message_bits = self.text_to_bits(hidden_message)
        length_bits = len(message_bits)
        length_header = format(length_bits, '016b')
        full_bits = length_header + message_bits
        k = self.get_k_bits()
        chunks, padding_bits = self.bits_to_chunks(full_bits, k)
        available_positions = self.get_available_positions_count(letter_count)
        if len(chunks) > available_positions:
            messagebox.showerror("Ошибка",
                                 f"Не хватает позиций для кодирования!\n"
                                 f"Нужно блоков: {len(chunks)}\n"
                                 f"Доступно позиций: {available_positions}")
            return
        russian_letters = []
        for paragraph in self.current_document.paragraphs:
            for char in paragraph.text:
                if 'А' <= char.upper() <= 'Я' or char.upper() == 'Ё':
                    russian_letters.append(char)
        encoding_positions = self.get_encoding_positions(len(russian_letters), len(chunks))
        self.hidden_message_positions = []
        encrypted_doc = Document()
        letter_index = 0
        chunk_index = 0
        for paragraph in self.current_document.paragraphs:
            new_paragraph = encrypted_doc.add_paragraph()
            for char in paragraph.text:
                run = new_paragraph.add_run(char)
                if 'А' <= char.upper() <= 'Я' or char.upper() == 'Ё':
                    if letter_index in encoding_positions and chunk_index < len(chunks):
                        current_chunk = chunks[chunk_index]
                        font_index = self.chunk_to_index(current_chunk)
                        font_index = font_index % self.num_fonts
                        if font_index < len(self.current_fonts):
                            font_name, font_size = self.current_fonts[font_index]
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            self.hidden_message_positions.append({
                                'position': letter_index,
                                'letter': char,
                                'chunk_bits': current_chunk,
                                'font_index': font_index,
                                'font': font_name,
                                'size': font_size
                            })
                        else:
                            run.font.name = 'Arial'
                            run.font.size = Pt(self.base_font_size)
                        chunk_index += 1
                    else:
                        run.font.name = 'Arial'
                        run.font.size = Pt(self.base_font_size)
                    letter_index += 1
                else:
                    run.font.name = 'Arial'
                    run.font.size = Pt(self.base_font_size)
        self.current_document = encrypted_doc
        new_text = ""
        for paragraph in encrypted_doc.paragraphs:
            new_text += paragraph.text + "\n"
        self.text_display.delete(1.0, END)
        self.text_display.insert(1.0, new_text)
        self.update_statistics()
        messagebox.showinfo("Успех",
                            f"Сообщение успешно зашифровано!\n\n"
                            f"Длина сообщения: {len(hidden_message)} символов\n"
                            f"Битов сообщения: {len(message_bits)}\n"
                            f"Использовано блоков: {len(chunks)}\n"
                            f"Ёмкость канала: {self.get_capacity():.2f} бит/символ (k={k} бит)\n"
                            f"Использовано шрифтов: {self.num_fonts}\n"
                            f"Padding бит: {padding_bits}")

    def decrypt_message(self):
        if not self.current_document:
            messagebox.showwarning("Предупреждение", "Сначала откройте документ!")
            return
        k = self.get_k_bits()
        extracted_indices = []
        self.hidden_message_positions = []
        position = 0
        for paragraph in self.current_document.paragraphs:
            for run in paragraph.runs:
                text = run.text
                font_name = run.font.name
                font_size = run.font.size.pt if run.font.size else None
                for char in text:
                    if 'А' <= char.upper() <= 'Я' or char.upper() == 'Ё':
                        if font_name and font_size:
                            key = f"{font_name}_{int(font_size)}"
                            if key in self.reverse_font_mapping:
                                font_index = self.reverse_font_mapping[key]
                                extracted_indices.append(font_index)
                                self.hidden_message_positions.append({
                                    'position': position,
                                    'letter': char,
                                    'font_index': font_index,
                                    'font': font_name,
                                    'size': int(font_size)
                                })
                        position += 1
        if not extracted_indices:
            self.result_display.delete(1.0, END)
            self.result_display.insert(1.0, "Скрытое сообщение не найдено!")
            messagebox.showinfo("Информация", "Скрытое сообщение не найдено!")
            return
        extracted_bits = ""
        for index in extracted_indices:
            extracted_bits += self.index_to_bits(index, k)
        if len(extracted_bits) < 16:
            self.result_display.delete(1.0, END)
            self.result_display.insert(1.0, "Ошибка: недостаточно данных для декодирования!")
            messagebox.showerror("Ошибка", "Недостаточно данных для декодирования!")
            return
        length_bits = extracted_bits[:16]
        message_length = int(length_bits, 2)
        message_bits = extracted_bits[16:16 + message_length]
        if len(message_bits) < message_length:
            self.result_display.delete(1.0, END)
            self.result_display.insert(1.0, f"Ошибка: ожидалось {message_length} бит, получено {len(message_bits)}")
            messagebox.showerror("Ошиб", f"Недостаточно бит для декодирования!")
            return
        hidden_message = self.bits_to_text(message_bits)
        self.result_display.delete(1.0, END)
        if hidden_message:
            self.result_display.insert(1.0, hidden_message)
            messagebox.showinfo("Успех",
                                f"Сообщение успешно расшифровано!\n"
                                f"Длина: {len(hidden_message)} символов\n"
                                f"Извлечено бит: {len(message_bits)}\n"
                                f"k = {k} бит на символ")
        else:
            self.result_display.insert(1.0, "Ошибка декодирования!")
            messagebox.showinfo("Информация", "Не удалось декодировать сообщение!")

    def show_encryption_map(self):
        if not self.hidden_message_positions:
            messagebox.showinfo("Информация",
                                "Нет данных о скрытом сообщении.\n"
                                "Сначала зашифруйте или расшифруйте сообщение.")
            return
        map_window = Toplevel(self.root)
        map_window.title("Карта скрытия сообщения")
        map_window.geometry("850x500")
        Label(map_window, text="Позиции скрытого сообщения",
              font=("Arial", 14, "bold")).pack(pady=10)
        text_frame = Frame(map_window)
        text_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        map_text = scrolledtext.ScrolledText(text_frame, wrap=WORD,
                                             font=("Courier", 10))
        map_text.pack(fill=BOTH, expand=True)
        has_bits = 'chunk_bits' in self.hidden_message_positions[0] if self.hidden_message_positions else False
        if has_bits:
            map_text.insert(END, "Позиция | Исх.буква | Биты      | Индекс | Шрифт                   | Размер\n")
            map_text.insert(END, "-" * 100 + "\n")
            for item in self.hidden_message_positions:
                map_text.insert(END,
                                f"{item['position']:7d} | {item['letter']:9s} | "
                                f"{item.get('chunk_bits', '?'):9s} | {item['font_index']:6d} | {item['font']:23s} | {item['size']} pt\n")
        else:
            map_text.insert(END, "Позиция | Исх.буква | Индекс | Шрифт                   | Размер\n")
            map_text.insert(END, "-" * 80 + "\n")
            for item in self.hidden_message_positions:
                map_text.insert(END,
                                f"{item['position']:7d} | {item['letter']:9s} | "
                                f"{item['font_index']:6d} | {item['font']:23s} | {item['size']} pt\n")
        map_text.insert(END, "\n" + "=" * 100 + "\n")
        map_text.insert(END, f"Всего позиций: {len(self.hidden_message_positions)}\n")
        map_text.insert(END, f"Ёмкость канала: {self.get_capacity():.2f} бит/символ\n")
        map_text.insert(END, f"k = {self.get_k_bits()} бит на символ\n")
        map_text.insert(END, f"Используется шрифтов: {self.num_fonts}\n")
        Button(map_window, text="Закрыть",
               command=map_window.destroy,
               bg="#607D8B", fg="white",
               padx=20, pady=5).pack(pady=10)

    def clear_encryption(self):
        if not self.current_document:
            messagebox.showwarning("Предупреждение", "Сначала откройте документ!")
            return
        answer = messagebox.askyesno("Подтверждение",
                                     "Вы уверены, что хотите очистить текст от шифрования?\n"
                                     "Все шрифты будут заменены на Arial.")
        if not answer:
            return
        try:
            cleared_doc = Document()
            for paragraph in self.current_document.paragraphs:
                new_paragraph = cleared_doc.add_paragraph()
                run = new_paragraph.add_run(paragraph.text)
                run.font.name = 'Arial'
                run.font.size = Pt(self.base_font_size)
            self.current_document = cleared_doc
            self.hidden_message_positions = []
            new_text = ""
            for paragraph in cleared_doc.paragraphs:
                new_text += paragraph.text + "\n"
            self.text_display.delete(1.0, END)
            self.text_display.insert(1.0, new_text)
            self.result_display.delete(1.0, END)
            self.update_statistics()
            messagebox.showinfo("Успех", "Текст успешно очищен от шифрования!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось очистить текст: {str(e)}")

    def save_decryption_result(self):
        result_text = self.result_display.get(1.0, END).strip()
        if not result_text:
            messagebox.showwarning("Предупреждение", "Нет результата для сохранения!")
            return
        try:
            file_path = filedialog.asksaveasfilename(
                title="Сохранить расшифровку",
                defaultextension=".docx",
                filetypes=[("DOCX файлы", "*.docx"), ("Текстовые файлы", "*.txt")]
            )
            if not file_path:
                return
            if file_path.endswith('.txt'):
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(result_text)
            else:
                result_doc = Document()
                paragraph = result_doc.add_paragraph()
                run = paragraph.add_run(result_text)
                run.font.name = 'Arial'
                run.font.size = Pt(self.base_font_size)
                result_doc.save(file_path)
            messagebox.showinfo("Успех", "Результат успешно сохранен!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить результат: {str(e)}")

    def on_text_change(self, event=None):
        self.update_capacity_indicator()


def main():
    root = Tk()
    app = SteganographyApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()